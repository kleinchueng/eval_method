import json

import docx
from transformers import AutoTokenizer, AutoModelForCausalLM

from builtins import PendingDeprecationWarning

from utils import convert_to_json
from metric.evaluator import get_evaluator


def compute_the_eval_score(file_name):
    # Example for dialogue response generation
    task = 'dialogue'

    # a list of dialogue histories
    src_list = []
    # a list of additional context that should be included into the generated response
    context_list = []
    # a list of model outputs to be evaluated
    output_list = []
    # Load the dataset
    with open(file_name, "r", encoding="utf-8") as f:
        data = json.load(f)

    for item in data:
        prompt = item["src"]
        src_list.append(prompt)
        generated = item["generated"]
        output_list.append(generated)
        context = item["context"]
        context_list.append(context)
    data = convert_to_json(output_list=output_list,
                           src_list=src_list, context_list=context_list)

    # Initialize evaluator for a specific task
    evaluator = get_evaluator(task)
    # Get multi-dimensional evaluation scores
    eval_scores = evaluator.evaluate(data, print_result=True)
    print(eval_scores)
    return eval_scores

file1 = "F:\dataset\\test1.json"
file2 = "F:\dataset\chatglm.json"
v_scores = compute_the_eval_score(file1)
chatglm_scores = compute_the_eval_score(file2)

document = docx.Document()

with open(file1, "r", encoding="utf-8") as f:
    data = json.load(f)
with open(file2, "r", encoding="utf-8") as f:
    data2 = json.load(f)

dims = list(v_scores[0].keys())
j = 0
while j < len(v_scores):
    document.add_paragraph('Q'+str(j+1)+'&A and eval_score are shown below:')
    table1 = document.add_table(rows=1, cols=3)
    table1.style = 'Table Grid'
    hdr_cells1 = table1.rows[0].cells
    hdr_cells1[0].text = '问题'
    hdr_cells1[1].text = 'Vicuna_Answer'
    hdr_cells1[2].text = 'ChatGLM_Answer'
    row_cells = table1.add_row().cells
    row_cells[0].text = data[j]["src"] + data[j]["context"]
    row_cells[1].text = data[j]["generated"]
    row_cells[2].text = data2[j]["generated"]
    table = document.add_table(rows=1, cols=3)
    table.style = 'Table Grid'

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Dimensions'
    hdr_cells[1].text = 'Vicuna_Score'
    hdr_cells[2].text = 'ChatGLM_Score'
    for dim in dims:
        row_cells = table.add_row().cells
        row_cells[0].text = dim
        row_cells[1].text = str(round(v_scores[j][dim]*100, 2))
        row_cells[2].text = str(round(chatglm_scores[j][dim]*100, 2))
    j = j + 1
    document.add_paragraph(" ")
# dims2 = list(chatglm_scores[0].keys())
# for dim in dims:
#     cur_score = 0
#     cur_score1 = 0
#     for i in range(len(v_scores)):
#         cur_score += v_scores[i][dim]
#         cur_score1 += chatglm_scores[i][dim]
#     row_cells = table.add_row().cells
#     row_cells[0].text = dim
#     row_cells[1].text = str(round(cur_score / len(v_scores), 6))
#     row_cells[2].text = str(round(cur_score / len(v_scores), 6))
# document.add_paragraph(table)

# i = 0
# while i < len(data):
#     row_cells = table1.add_row().cells
#     row_cells[0].text = data[i]["src"]+data[i]["context"]
#     row_cells[1].text = data[i]["generated"]
#     row_cells[2].text = data2[i]["generated"]
#     i = i + 1

document.save('table.docx')
