import os
import json
import docx
from prettytable import PrettyTable

document = docx.Document()
ques = []
with open('F:\毕设模型评估\FastChat\\fastchat\eval\question.jsonl', "r", encoding="utf-8") as f:
    for line in f:
        tmp = json.loads(line)
        ques.append(tmp["text"])

score = []
vicuna_score = []
chatglm_score = []
txt = []

with open('F:\毕设模型评估\FastChat\\fastchat\eval\\review-output——t.jsonl', "r", encoding="gbk") as f:
    for line in f:
        data = json.loads(line)
        vicuna_score.append(data["score"][0])
        chatglm_score.append(data["score"][1])
        text = data["text"]
        split_text = text.split('\n\n')
        score.append(split_text[0])
        txt.append(split_text[1])

table = document.add_table(rows=1, cols=4)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '问题'
hdr_cells[1].text = 'Vicuna_Score'
hdr_cells[2].text = 'ChatGLM_Score'
hdr_cells[3].text = '评价理由'

# table = PrettyTable(['question', 'vicuna_ans', 'chatglm_ans', 'vicuna_Score', 'chatglm_Score'])

i = 0
# 将socre里的每一项用\n分割，将将分割的小项用空格分割，然后将分割的小项分别赋值给score1_list和score2_list
print('\nEvaluation scores are shown below:')
while i < len(score):
    split_text = text.split('\n')
    score1 = split_text[1]
    score2 = split_text[2]
    score1_list = score1.split(' ')
    score2_list = score2.split(' ')

    text1 = "可用性:" + score1_list[0] + "\n" + "可读性:" + score1_list[1] + "\n" + "可理解性:" + score1_list[
        2] + "\n" + "可信度:" + score1_list[3] + "\n" + "总分：" + str(vicuna_score[i])
    text2 = "可用性:" + score2_list[0] + "\n" + "可读性:" + score2_list[1] + "\n" + "可理解性:" + score2_list[
        2] + "\n" + "可信度:" + score2_list[3] + "\n" + "总分：" + str(chatglm_score[i])
    #添加到table
    row_cells = table.add_row().cells
    row_cells[0].text = ques[i]
    row_cells[1].text = str(text1)
    row_cells[2].text = str(text2)
    row_cells[3].text = txt[i]
    i = i + 1
#修改document中的字体为宋体5号字
for paragraph in document.paragraphs:
    paragraph.style.font.name = '宋体'
    paragraph.style.font.size = docx.shared.Pt(5)
document.save('table1.docx')
